from docx import Document


def generate_news_script(title, content):
    return f"""【校园新闻播报稿】

各位老师、同学们，大家好。

下面为大家播报一则校园新闻。

新闻标题：

{title}

新闻内容：

{content}

以上就是本条新闻的全部内容。
感谢大家收听。
"""


def save_to_word(text):
    doc = Document()

    doc.add_heading('校园新闻播报稿', level=1)

    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc.save("output.docx")


def main():
    title = input("请输入新闻标题：")
    content = input("请输入新闻正文：")

    result = generate_news_script(title, content)

    save_to_word(result)

    print("生成成功")
    print("已生成 output.docx")


if __name__ == "__main__":
    main()
